from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.spinner import Spinner
from premises_calculator.room import Room
from premises_calculator.apartment import Apartment
from premises_calculator.building import Building
from docx import Document
import os

class PremisesApp(App):
    def build(self):
        self.results = []
        layout = BoxLayout(orientation='vertical', padding=10, spacing=10)

        # Выбор типа помещения
        layout.add_widget(Label(text="Тип помещения:"))
        self.premise_type = Spinner(
            text='Комната',
            values=('Комната', 'Квартира', 'Многоэтажный дом')
        )
        layout.add_widget(self.premise_type)

        # Поля ввода для комнаты
        self.room_layout = BoxLayout(spacing=5)
        self.room_layout.add_widget(Label(text="Длина (м):"))
        self.length_input = TextInput(multiline=False)
        self.room_layout.add_widget(self.length_input)
        self.room_layout.add_widget(Label(text="Ширина (м):"))
        self.width_input = TextInput(multiline=False)
        self.room_layout.add_widget(self.width_input)
        self.room_layout.add_widget(Label(text="Высота (м):"))
        self.height_input = TextInput(multiline=False)
        self.room_layout.add_widget(self.height_input)
        layout.add_widget(self.room_layout)

        # Поля ввода для квартиры
        self.apartment_layout = BoxLayout(spacing=5)
        self.apartment_layout.add_widget(Label(text="Количество комнат:"))
        self.room_count_input = TextInput(multiline=False)
        self.apartment_layout.add_widget(self.room_count_input)
        layout.add_widget(self.apartment_layout)

        # Поля ввода для дома
        self.building_layout = BoxLayout(spacing=5)
        self.building_layout.add_widget(Label(text="Количество этажей:"))
        self.floor_count_input = TextInput(multiline=False)
        self.building_layout.add_widget(self.floor_count_input)
        self.building_layout.add_widget(Label(text="Квартир на этаже:"))
        self.apt_per_floor_input = TextInput(multiline=False)
        self.building_layout.add_widget(self.apt_per_floor_input)
        layout.add_widget(self.building_layout)

        # Кнопки
        calculate_button = Button(text="Рассчитать")
        calculate_button.bind(on_press=self.calculate)
        layout.add_widget(calculate_button)

        save_button = Button(text="Сохранить отчёт")
        save_button.bind(on_press=self.save_report)
        layout.add_widget(save_button)

        # Поле вывода результатов
        self.result_label = Label(text="", size_hint=(1, 0.3))
        layout.add_widget(self.result_label)

        return layout

    def calculate(self, instance):
        try:
            self.results = []
            premise_type = self.premise_type.text

            if premise_type == "Комната":
                length = float(self.length_input.text)
                width = float(self.width_input.text)
                height = float(self.height_input.text)
                room = Room(length, width, height)
                self.results.append(str(room))

            elif premise_type == "Квартира":
                room_count = int(self.room_count_input.text)
                apartment = Apartment()
                length = float(self.length_input.text)
                width = float(self.width_input.text)
                height = float(self.height_input.text)
                for _ in range(room_count):
                    apartment.add_room(length, width, height)
                self.results.append(str(apartment))

            elif premise_type == "Многоэтажный дом":
                floor_count = int(self.floor_count_input.text)
                apt_per_floor = int(self.apt_per_floor_input.text)
                building = Building()
                length = float(self.length_input.text)
                width = float(self.width_input.text)
                height = float(self.height_input.text)
                for _ in range(floor_count * apt_per_floor):
                    apartment = Apartment()
                    apartment.add_room(length, width, height)
                    building.add_apartment(apartment)
                self.results.append(str(building))

            self.result_label.text = "\n".join(self.results)

        except ValueError as e:
            self.result_label.text = f"Ошибка: {str(e)} или некорректный ввод"

    def save_report(self, instance):
        if not self.results:
            self.result_label.text = "Сначала выполните расчёт"
            return

        doc = Document()
        doc.add_heading("Отчёт по расчёту помещений", 0)
        for result in self.results:
            doc.add_paragraph(result)

        report_path = "premises_report.docx"
        doc.save(report_path)
        self.result_label.text = f"Отчёт сохранён как {report_path}"

if __name__ == "__main__":
    PremisesApp().run()