import os
from datetime import datetime
from typing import Dict, Any

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

class ReportGenerator:
    """Класс для генерации отчетов"""
    
    def __init__(self):
        self.report_data = {}
        
    def set_data(self, calculation_data: Dict, thermal_data: Dict, 
                 cost_data: Dict = None):
        """Установка данных для отчета"""
        self.report_data = {
            'calculation': calculation_data,
            'thermal': thermal_data,
            'cost': cost_data,
            'timestamp': datetime.now().strftime("%d.%m.%Y %H:%M")
        }
    
    def generate_docx_report(self, filename: str = None) -> str:
        """
        Генерация отчета в формате DOCX
        Returns:
            Путь к созданному файлу
        """
        if not DOCX_AVAILABLE:
            raise ImportError("Пакет python-docx не установлен")
            
        if filename is None:
            filename = f"heating_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('Отчет по расчету отопления помещений', 0)
        title.alignment = 1  # Центровка
        
        # Дата создания
        doc.add_paragraph(f'Дата создания отчета: {self.report_data["timestamp"]}')
        doc.add_paragraph()
        
        # Расчет площади
        doc.add_heading('1. Расчет площади помещения', level=1)
        calc_data = self.report_data['calculation']
        
        p = doc.add_paragraph()
        p.add_run(f'Тип помещения: ').bold = True
        p.add_run(calc_data['type'])
        
        if 'dimensions' in calc_data:
            p = doc.add_paragraph()
            p.add_run(f'Размеры: ').bold = True
            p.add_run(f"{calc_data['dimensions']} м")
        
        p = doc.add_paragraph()
        p.add_run(f'Площадь пола: ').bold = True
        p.add_run(f"{calc_data['floor_area']} м²")
        
        if 'wall_area' in calc_data:
            p = doc.add_paragraph()
            p.add_run(f'Площадь стен: ').bold = True
            p.add_run(f"{calc_data['wall_area']} м²")
        
        p = doc.add_paragraph()
        p.add_run(f'Общая площадь: ').bold = True
        p.add_run(f"{calc_data['total_area']} м²")
        
        p = doc.add_paragraph()
        p.add_run(f'Объем: ').bold = True
        p.add_run(f"{calc_data['volume']} м³")
        
        # Тепловой расчет
        doc.add_heading('2. Расчет тепловой мощности', level=1)
        thermal_data = self.report_data['thermal']
        
        p = doc.add_paragraph()
        p.add_run(f'Базовая мощность: ').bold = True
        p.add_run(f"{thermal_data['base_power']} Вт")
        
        p = doc.add_paragraph()
        p.add_run(f'Скорректированная мощность: ').bold = True
        p.add_run(f"{thermal_data['adjusted_power']} Вт")
        
        p = doc.add_paragraph()
        p.add_run(f'Рекомендуемая мощность: ').bold = True
        p.add_run(f"{thermal_data['recommended_power']} Вт ({thermal_data['recommended_power']/1000:.1f} кВт)")
        
        p = doc.add_paragraph()
        p.add_run(f'Удельная мощность: ').bold = True
        p.add_run(f"{thermal_data['power_per_m2']} Вт/м²")
        
        # Коэффициенты
        doc.add_heading('3. Коэффициенты расчета', level=1)
        
        p = doc.add_paragraph()
        p.add_run(f'Коэффициент климата: ').bold = True
        p.add_run(f"{thermal_data['climate_coeff']}")
        
        p = doc.add_paragraph()
        p.add_run(f'Коэффициент теплоизоляции: ').bold = True
        p.add_run(f"{thermal_data['insulation_coeff']}")
        
        p = doc.add_paragraph()
        p.add_run(f'Коэффициент высоты: ').bold = True
        p.add_run(f"{thermal_data['height_coeff']}")
        
        # Стоимость (если есть)
        if self.report_data['cost']:
            doc.add_heading('4. Расчет стоимости отопления', level=1)
            cost_data = self.report_data['cost']
            
            p = doc.add_paragraph()
            p.add_run(f'Мощность оборудования: ').bold = True
            p.add_run(f"{cost_data['power_kw']} кВт")
            
            p = doc.add_paragraph()
            p.add_run(f'Месячное потребление: ').bold = True
            p.add_run(f"{cost_data['monthly_consumption']} кВт·ч")
            
            p = doc.add_paragraph()
            p.add_run(f'Стоимость в месяц: ').bold = True
            p.add_run(f"{cost_data['monthly_cost']} руб.")
            
            p = doc.add_paragraph()
            p.add_run(f'Стоимость за сезон: ').bold = True
            p.add_run(f"{cost_data['yearly_cost']} руб.")
        
        doc.save(filename)
        return os.path.abspath(filename)
    
    def generate_xlsx_report(self, filename: str = None) -> str:
        """
        Генерация отчета в формате XLSX
        Returns:
            Путь к созданному файлу
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError("Пакет openpyxl не установлен")
            
        if filename is None:
            filename = f"heating_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Расчет отопления"
        
        # Стили
        header_font = Font(bold=True, size=14)
        subheader_font = Font(bold=True, size=12)
        bold_font = Font(bold=True)
        center_alignment = Alignment(horizontal='center')
        
        # Заголовок
        ws['A1'] = 'ОТЧЕТ ПО РАСЧЕТУ ОТОПЛЕНИЯ ПОМЕЩЕНИЙ'
        ws['A1'].font = header_font
        ws['A1'].alignment = center_alignment
        ws.merge_cells('A1:D1')
        
        ws['A2'] = f'Дата: {self.report_data["timestamp"]}'
        ws['A2'].font = bold_font
        
        row = 4
        
        # Расчет площади
        ws[f'A{row}'] = 'РАСЧЕТ ПЛОЩАДИ ПОМЕЩЕНИЯ'
        ws[f'A{row}'].font = subheader_font
        row += 1
        
        calc_data = self.report_data['calculation']
        
        ws[f'A{row}'] = 'Параметр'
        ws[f'B{row}'] = 'Значение'
        ws[f'C{row}'] = 'Единица'
        for col in ['A', 'B', 'C']:
            ws[f'{col}{row}'].font = bold_font
        row += 1
        
        ws[f'A{row}'] = 'Тип помещения'
        ws[f'B{row}'] = calc_data['type']
        row += 1
        
        if 'dimensions' in calc_data:
            ws[f'A{row}'] = 'Размеры'
            ws[f'B{row}'] = calc_data['dimensions']
            ws[f'C{row}'] = 'м'
            row += 1
        
        ws[f'A{row}'] = 'Площадь пола'
        ws[f'B{row}'] = calc_data['floor_area']
        ws[f'C{row}'] = 'м²'
        row += 1
        
        if 'wall_area' in calc_data:
            ws[f'A{row}'] = 'Площадь стен'
            ws[f'B{row}'] = calc_data['wall_area']
            ws[f'C{row}'] = 'м²'
            row += 1
        
        ws[f'A{row}'] = 'Общая площадь'
        ws[f'B{row}'] = calc_data['total_area']
        ws[f'C{row}'] = 'м²'
        row += 1
        
        ws[f'A{row}'] = 'Объем'
        ws[f'B{row}'] = calc_data['volume']
        ws[f'C{row}'] = 'м³'
        row += 2
        
        # Тепловой расчет
        ws[f'A{row}'] = 'РАСЧЕТ ТЕПЛОВОЙ МОЩНОСТИ'
        ws[f'A{row}'].font = subheader_font
        row += 1
        
        thermal_data = self.report_data['thermal']
        
        ws[f'A{row}'] = 'Параметр'
        ws[f'B{row}'] = 'Значение'
        ws[f'C{row}'] = 'Единица'
        for col in ['A', 'B', 'C']:
            ws[f'{col}{row}'].font = bold_font
        row += 1
        
        ws[f'A{row}'] = 'Базовая мощность'
        ws[f'B{row}'] = thermal_data['base_power']
        ws[f'C{row}'] = 'Вт'
        row += 1
        
        ws[f'A{row}'] = 'Скорректированная мощность'
        ws[f'B{row}'] = thermal_data['adjusted_power']
        ws[f'C{row}'] = 'Вт'
        row += 1
        
        ws[f'A{row}'] = 'Рекомендуемая мощность'
        ws[f'B{row}'] = thermal_data['recommended_power']
        ws[f'C{row}'] = 'Вт'
        row += 1
        
        ws[f'A{row}'] = 'Удельная мощность'
        ws[f'B{row}'] = thermal_data['power_per_m2']
        ws[f'C{row}'] = 'Вт/м²'
        row += 2
        
        # Коэффициенты
        ws[f'A{row}'] = 'КОЭФФИЦИЕНТЫ РАСЧЕТА'
        ws[f'A{row}'].font = subheader_font
        row += 1
        
        coeffs = [
            ('Коэффициент климата', thermal_data['climate_coeff']),
            ('Коэффициент теплоизоляции', thermal_data['insulation_coeff']),
            ('Коэффициент высоты', thermal_data['height_coeff'])
        ]
        
        for name, value in coeffs:
            ws[f'A{row}'] = name
            ws[f'B{row}'] = value
            row += 1
        
        # Стоимость (если есть)
        if self.report_data['cost']:
            row += 1
            ws[f'A{row}'] = 'РАСЧЕТ СТОИМОСТИ ОТОПЛЕНИЯ'
            ws[f'A{row}'].font = subheader_font
            row += 1
            
            cost_data = self.report_data['cost']
            costs = [
                ('Мощность оборудования', cost_data['power_kw'], 'кВт'),
                ('Месячное потребление', cost_data['monthly_consumption'], 'кВт·ч'),
                ('Стоимость в месяц', cost_data['monthly_cost'], 'руб.'),
                ('Стоимость за сезон', cost_data['yearly_cost'], 'руб.')
            ]
            
            for name, value, unit in costs:
                ws[f'A{row}'] = name
                ws[f'B{row}'] = value
                ws[f'C{row}'] = unit
                row += 1
        
        # Автоподгон ширины колонок
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(filename)
        return os.path.abspath(filename)